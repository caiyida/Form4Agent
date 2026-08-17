"""Behavior tests for XML-safe Word placeholder replacement."""

from io import BytesIO
from pathlib import Path
import sys
import unittest
from zipfile import ZipFile

from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT / "src"))

from form4_engine import fill_form  # noqa: E402
from word_helper import find_placeholders, replace_placeholder  # noqa: E402


class PlaceholderReplacementTests(unittest.TestCase):
    def test_replaces_placeholder_split_across_runs(self):
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.add_run("Before {{tenant")
        paragraph.add_run("1_name}} after")

        count = replace_placeholder(document, "tenant1_name", "Test User")

        self.assertEqual(count, 1)
        self.assertEqual(paragraph.text, "Before Test User after")
        self.assertEqual(find_placeholders(document), set())

    def test_replaces_repeated_placeholder_with_longer_value(self):
        document = Document()
        paragraph = document.add_paragraph("{{field}} / {{field}}")

        count = replace_placeholder(document, "field", "LONGER VALUE")

        self.assertEqual(count, 2)
        self.assertEqual(paragraph.text, "LONGER VALUE / LONGER VALUE")

    def test_generation_blanks_values_omitted_from_form_data(self):
        output = BytesIO()
        fill_form({}, output)
        output.seek(0)

        with ZipFile(output) as archive:
            xml_parts = [
                archive.read(name)
                for name in archive.namelist()
                if name.startswith("word/") and name.endswith(".xml")
            ]

        self.assertFalse(any(b"{{" in part or b"}}" in part for part in xml_parts))

    def test_generation_writes_requested_values(self):
        output = BytesIO()
        fill_form({"tenant1_name": "SYNTHETIC TENANT"}, output)
        output.seek(0)

        with ZipFile(output) as archive:
            document_xml = archive.read("word/document.xml")

        self.assertIn(b"SYNTHETIC TENANT", document_xml)

    def test_review_generation_removes_agent_marks_only(self):
        output = BytesIO()
        fill_form({}, output, include_agent_marks=False)
        output.seek(0)

        with ZipFile(output) as archive:
            document_xml = archive.read("word/document.xml")

        self.assertNotIn(b'descr="initials"', document_xml)
        self.assertNotIn(b'descr="signature"', document_xml)


if __name__ == "__main__":
    unittest.main()
